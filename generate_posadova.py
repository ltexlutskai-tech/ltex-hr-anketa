"""
Generate: Posadova Instruktsiia — Menedzher z prodazhu (Sales Manager)
Company: L-TEX (Secondopt + Bricabrac)
Output: output/posadova_menedzher_prodazhu.docx
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Helpers ──────────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Set margins inside a table cell (in twips)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tc_pr.append(margins)


def add_heading_styled(doc, text, level=1):
    """Add a heading with Arial font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_para(doc, text, bold=False, italic=False, size=12, alignment=None, space_after=Pt(4)):
    """Add a paragraph with consistent styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    return p


def add_numbered_item(doc, number, text, bold_prefix=""):
    """Add a numbered list item."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(1)
    if bold_prefix:
        run_num = p.add_run(f"{number}. {bold_prefix}")
        run_num.font.name = "Arial"
        run_num.font.size = Pt(12)
        run_num.bold = True
        run_text = p.add_run(f" {text}")
        run_text.font.name = "Arial"
        run_text.font.size = Pt(12)
    else:
        run = p.add_run(f"{number}. {text}")
        run.font.name = "Arial"
        run.font.size = Pt(12)
    return p


def add_bullet(doc, text, indent=1.5):
    """Add a bullet point item."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(indent)
    run_bullet = p.add_run("•  ")
    run_bullet.font.name = "Arial"
    run_bullet.font.size = Pt(12)
    run_text = p.add_run(text)
    run_text.font.name = "Arial"
    run_text.font.size = Pt(12)
    return p


def add_spacer(doc, height=6):
    """Add vertical spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(height)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("")
    run.font.size = Pt(1)


# ── Document Creation ────────────────────────────────────────────────

doc = Document()

# Page setup — A4 with Ukrainian standard margins
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

# Default font
style = doc.styles["Normal"]
font = style.font
font.name = "Arial"
font.size = Pt(12)

# Header
section = doc.sections[0]
header = section.header
header_para = header.paragraphs[0]
header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header_para.add_run("L-TEX  |  Посадова інструкція  |  Менеджер з продажу")
run.font.name = "Arial"
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ══════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════

add_spacer(doc, 12)

# Approval block — right aligned
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(2)
run = p.add_run("ЗАТВЕРДЖУЮ")
run.font.name = "Arial"
run.font.size = Pt(12)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(2)
run = p.add_run("Директор ТОВ «Л-ТЕКС»")
run.font.name = "Arial"
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(2)
run = p.add_run("_____________ / ________________ /")
run.font.name = "Arial"
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(6)
run = p.add_run("«____» ________________ 2026 р.")
run.font.name = "Arial"
run.font.size = Pt(12)

add_spacer(doc, 16)

# Main title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run("ПОСАДОВА ІНСТРУКЦІЯ")
run.font.name = "Arial"
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(4)
run = subtitle.add_run("МЕНЕДЖЕР З ПРОДАЖУ")
run.font.name = "Arial"
run.font.size = Pt(16)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

company_line = doc.add_paragraph()
company_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
company_line.paragraph_format.space_after = Pt(2)
run = company_line.add_run("ТОВ «Л-ТЕКС» (бренди Secondopt, Bricabrac)")
run.font.name = "Arial"
run.font.size = Pt(12)
run.italic = True

location_line = doc.add_paragraph()
location_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
location_line.paragraph_format.space_after = Pt(2)
run = location_line.add_run("м. Луцьк, Волинська обл.")
run.font.name = "Arial"
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_spacer(doc, 12)

# ══════════════════════════════════════════════════════════════════════
# 1. ЗАГАЛЬНІ ПОЛОЖЕННЯ
# ══════════════════════════════════════════════════════════════════════

add_heading_styled(doc, "1. Загальні положення", level=1)

general_provisions = [
    "Менеджер з продажу належить до категорії фахівців відділу продажів ТОВ «Л-ТЕКС».",
    "Призначення на посаду та звільнення з неї здійснюється наказом директора компанії.",
    "Менеджер з продажу безпосередньо підпорядковується керівнику компанії.",
    "Компанія L-TEX спеціалізується на оптовому продажу секонд-хенду та товарів Bric-a-Brac з поставками з Великобританії, Польщі, Канади, Німеччини та Італії.",
    "Менеджер працює з асортиментом 600+ позицій у категоріях: одяг, взуття, сумки, товари для дому, іграшки, Bric-a-Brac, косметика.",
    "Класифікація товару за сортами: Екстра, Крем, 1-й сорт, 2-й сорт, Сток.",
    "Усі ціни формуються та комунікуються клієнтам у валюті EUR.",
    "Мінімальне замовлення: від 10 кг.",
    "Команда відділу продажів: 5+ менеджерів з розмежуванням по територіях.",
    "У своїй діяльності менеджер з продажу керується чинним законодавством України, Статутом компанії, внутрішніми регламентами та цією посадовою інструкцією.",
]

for i, item in enumerate(general_provisions, 1):
    add_numbered_item(doc, f"1.{i}", item)

# ══════════════════════════════════════════════════════════════════════
# 2. ПІДПОРЯДКУВАННЯ ТА ВЗАЄМОДІЯ
# ══════════════════════════════════════════════════════════════════════

add_heading_styled(doc, "2. Підпорядкування та взаємодія", level=1)

add_para(doc, "Пряме підпорядкування:", bold=True, space_after=Pt(2))
add_bullet(doc, "Керівник компанії (директор ТОВ «Л-ТЕКС»)")

add_spacer(doc, 4)
add_para(doc, "Взаємодія:", bold=True, space_after=Pt(2))
add_bullet(doc, "Складська служба — координація відвантажень, наявність товару")
add_bullet(doc, "Бухгалтерія — контроль оплат, дебіторська заборгованість")
add_bullet(doc, "Інші менеджери з продажу — обмін досвідом, розмежування територій")
add_bullet(doc, "Маркетинг — зворотній зв'язок по попиту, участь у плануванні закупівель")

add_spacer(doc, 4)
add_para(doc, "Структура команди:", bold=True, space_after=Pt(2))
add_bullet(doc, "Команда: 5+ менеджерів з продажу")
add_bullet(doc, "Кожен менеджер отримує закріплену територію та готову базу клієнтів свого регіону")
add_bullet(doc, "Розмежування відповідальності по територіях між менеджерами")

# ══════════════════════════════════════════════════════════════════════
# 3. ПОСАДОВІ ОБОВ'ЯЗКИ
# ══════════════════════════════════════════════════════════════════════

add_heading_styled(doc, "3. Посадові обов'язки", level=1)

add_para(doc, "Менеджер з продажу зобов'язаний:", italic=True, space_after=Pt(4))

duties = [
    "Опрацьовувати вхідні заявки з усіх каналів (сайти secondopt.com.ua, bricabrac.com.ua, Telegram, Viber, телефон) — час відповіді до 1 години.",
    "Надсилати актуальний прайс-лист (600+ позицій у EUR) новим та потенційним клієнтам.",
    "Консультувати клієнтів по асортименту: категорії товарів, система сортів (Екстра, Крем, 1-й, 2-й сорт, Сток), ціноутворення у EUR.",
    "Надсилати відеоогляди товару з YouTube-каналу L-TEX для наочної демонстрації якості.",
    "Вести та актуалізувати клієнтську базу в 1С та Excel: контакти, історія замовлень, примітки.",
    "Фіксувати замовлення в системі обліку, контролювати своєчасність оплат.",
    "Координувати відвантаження зі складською службою: формування замовлень, терміни, логістика.",
    "Професійно працювати з запереченнями клієнтів, знаходити рішення та компроміси.",
    "Реактивувати «сплячих» клієнтів (відсутність покупок більше 60 днів): дзвінки, спецпропозиції, персональні умови.",
    "Контролювати дебіторську заборгованість закріплених клієнтів, вживати заходів для своєчасного погашення.",
    "Готувати та надавати керівнику щотижневий звіт: виручка, кількість нових клієнтів, стан воронки продажів.",
    "Брати участь у плануванні закупівель: надавати зворотній зв'язок по попиту, тенденціях ринку, побажаннях клієнтів.",
    "Утримувати існуючу базу клієнтів закріпленого регіону та активно залучати нових оптових покупців.",
    "Вивчати ринок, конкурентів, цінову політику для підвищення ефективності продажів.",
    "Дотримуватися стандартів комунікації компанії L-TEX у всіх каналах зв'язку.",
]

for i, duty in enumerate(duties, 1):
    add_numbered_item(doc, f"3.{i}", duty)

# ══════════════════════════════════════════════════════════════════════
# 4. ПРАВА
# ══════════════════════════════════════════════════════════════════════

add_heading_styled(doc, "4. Права", level=1)

add_para(doc, "Менеджер з продажу має право:", italic=True, space_after=Pt(4))

rights = [
    "Отримувати від керівництва інформацію, необхідну для виконання посадових обов'язків.",
    "Вносити пропозиції щодо вдосконалення роботи відділу продажів, асортименту, цінової політики.",
    "Приймати рішення в межах своєї компетенції щодо умов роботи з клієнтами (знижки, відстрочки — в рамках затверджених лімітів).",
    "Підвищувати кваліфікацію за рахунок компанії (тренінги з продажу, знання продукту).",
    "Отримувати корпоративний телефон та ноутбук для виконання службових обов'язків.",
    "Звертатися до керівника з питаннями, що виходять за межі компетенції.",
]

for i, right in enumerate(rights, 1):
    add_numbered_item(doc, f"4.{i}", right)

# ══════════════════════════════════════════════════════════════════════
# 5. ВІДПОВІДАЛЬНІСТЬ
# ══════════════════════════════════════════════════════════════════════

add_heading_styled(doc, "5. Відповідальність", level=1)

add_para(doc, "Менеджер з продажу несе відповідальність за:", italic=True, space_after=Pt(4))

responsibilities = [
    "Неналежне виконання або невиконання посадових обов'язків, передбачених цією інструкцією.",
    "Недосягнення встановлених KPI без об'єктивних причин.",
    "Втрату клієнтів з бази закріпленого регіону через неналежну роботу.",
    "Несвоєчасну обробку заявок (порушення нормативу відповіді — 1 година).",
    "Неконтрольовану дебіторську заборгованість закріплених клієнтів.",
    "Розголошення комерційної таємниці, клієнтської бази, цінових умов компанії.",
    "Порушення трудової дисципліни та внутрішніх регламентів компанії.",
    "Завдання матеріальної шкоди компанії — в межах, визначених чинним законодавством України.",
]

for i, resp in enumerate(responsibilities, 1):
    add_numbered_item(doc, f"5.{i}", resp)

# ══════════════════════════════════════════════════════════════════════
# 6. KPI ТА ПОКАЗНИКИ ЕФЕКТИВНОСТІ
# ══════════════════════════════════════════════════════════════════════

add_heading_styled(doc, "6. KPI та показники ефективності", level=1)

add_para(doc, "Ефективність роботи менеджера з продажу оцінюється за такими показниками:",
         italic=True, space_after=Pt(6))

# KPI Table
table = doc.add_table(rows=7, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(7)
    row.cells[1].width = Cm(5.5)
    row.cells[2].width = Cm(4)

# Header row
headers = ["Показник", "Норма", "Одиниця виміру"]
for i, header_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, "1A1A2E")
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header_text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# KPI data
kpi_data = [
    ("Виручка", "Місячний план (встановлює керівник)", "EUR"),
    ("Нові клієнти", "Мінімум 10", "клієнтів / місяць"),
    ("Конверсія вхідних заявок", "Мінімум 40%", "%"),
    ("Середній чек замовлення", "Від 500", "EUR"),
    ("Час відповіді на заявку", "До 1 години", "година"),
    ("Частка повторних покупок", "Мінімум 60% бази", "%"),
]

for row_idx, (metric, norm, unit) in enumerate(kpi_data, 1):
    row = table.rows[row_idx]

    # Alternate row shading
    bg = "F5F5FA" if row_idx % 2 == 0 else "FFFFFF"

    for col_idx, value in enumerate([metric, norm, unit]):
        cell = row.cells[col_idx]
        set_cell_shading(cell, bg)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        if col_idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        if col_idx == 0:
            run.bold = True

add_spacer(doc, 6)
add_para(doc, "Примітка: конкретний місячний план виручки встановлюється керівником індивідуально та може переглядатися щоквартально.",
         italic=True, size=10, space_after=Pt(4))

# ══════════════════════════════════════════════════════════════════════
# 7. ВИМОГИ ДО КАНДИДАТА
# ══════════════════════════════════════════════════════════════════════

add_heading_styled(doc, "7. Вимоги до кандидата", level=1)

add_heading_styled(doc, "7.1. Обов'язкові вимоги", level=2)

mandatory = [
    "Грамотна усна та письмова українська мова.",
    "Впевнений користувач: смартфон, месенджери (Telegram, Viber), Excel.",
    "Відповідальність, комунікабельність, орієнтація на результат.",
    "Готовність до відряджень у межах закріпленої території.",
    "Вища або незакінчена вища освіта.",
]

for i, req in enumerate(mandatory, 1):
    add_numbered_item(doc, f"7.1.{i}", req)

add_heading_styled(doc, "7.2. Переваги (буде плюсом)", level=2)

preferred = [
    "Досвід у продажах або роботі з клієнтами (B2B — значна перевага).",
    "Водійські права категорії В, наявність власного авто.",
    "Досвід роботи з 1С.",
    "Знання ринку секонд-хенду або оптової торгівлі.",
    "Знання англійської мови на рівні читання (для роботи з іноземними постачальниками).",
]

for i, pref in enumerate(preferred, 1):
    add_numbered_item(doc, f"7.2.{i}", pref)

# ══════════════════════════════════════════════════════════════════════
# 8. УМОВИ ПРАЦІ
# ══════════════════════════════════════════════════════════════════════

add_heading_styled(doc, "8. Умови праці", level=1)

add_heading_styled(doc, "8.1. Загальні умови", level=2)
add_bullet(doc, "Формат роботи: офіс, м. Луцьк")
add_bullet(doc, "Графік: понеділок — п'ятниця, 09:00 — 18:00")
add_bullet(doc, "Забезпечення: корпоративний телефон + ноутбук")
add_bullet(doc, "Робочі інструменти: 1С, Excel, Telegram, Viber")

add_heading_styled(doc, "8.2. Мотивація", level=2)
add_bullet(doc, "Система оплати: ставка + відсоток від продажів + бонуси за виконання KPI")
add_bullet(doc, "Конкретні розміри ставки та відсотку погоджуються індивідуально")

add_heading_styled(doc, "8.3. Випробування та адаптація", level=2)
add_bullet(doc, "Випробувальний термін: 1 місяць")
add_bullet(doc, "Повне входження в посаду: 3 місяці")
add_bullet(doc, "Менеджер отримує готову базу клієнтів свого регіону з першого дня роботи")
add_bullet(doc, "Якщо після 3 місяців план систематично не виконується — перегляд мотивації та подальшої співпраці")

# ══════════════════════════════════════════════════════════════════════
# 9. ПІДПИСИ
# ══════════════════════════════════════════════════════════════════════

add_heading_styled(doc, "9. Підписи", level=1)

add_spacer(doc, 12)

# Signature table (no borders)
sig_table = doc.add_table(rows=4, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for row in sig_table.rows:
    for cell in row.cells:
        cell.width = Cm(8)
        # Remove borders
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for side in ["top", "bottom", "left", "right"]:
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "none")
            border.set(qn("w:sz"), "0")
            borders.append(border)
        tc_pr.append(borders)

# Row 0: titles
for col, title_text in enumerate(["ДИРЕКТОР", "ПРАЦІВНИК"]):
    p = sig_table.rows[0].cells[col].paragraphs[0]
    run = p.add_run(title_text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = True

# Row 1: company/position
p = sig_table.rows[1].cells[0].paragraphs[0]
run = p.add_run("ТОВ «Л-ТЕКС»")
run.font.name = "Arial"
run.font.size = Pt(11)

p = sig_table.rows[1].cells[1].paragraphs[0]
run = p.add_run("Менеджер з продажу")
run.font.name = "Arial"
run.font.size = Pt(11)

# Row 2: signature lines
for col in range(2):
    p = sig_table.rows[2].cells[col].paragraphs[0]
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("_____________ / ________________ /")
    run.font.name = "Arial"
    run.font.size = Pt(11)

# Row 3: date
for col in range(2):
    p = sig_table.rows[3].cells[col].paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("«____» ________________ 2026 р.")
    run.font.name = "Arial"
    run.font.size = Pt(11)

add_spacer(doc, 12)

# Final note
add_para(doc, "З посадовою інструкцією ознайомлений(а), один примірник отримав(ла):",
         italic=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_spacer(doc, 8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("_____________ / ________________ /          «____» ________________ 2026 р.")
run.font.name = "Arial"
run.font.size = Pt(10)

# ── Save ─────────────────────────────────────────────────────────────

output_dir = os.path.join(os.path.dirname(__file__), "output")
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "posadova_menedzher_prodazhu.docx")
doc.save(output_path)

file_size = os.path.getsize(output_path)
print(f"Document saved: {output_path}")
print(f"File size: {file_size:,} bytes ({file_size / 1024:.1f} KB)")
print(f"Sections: {len(doc.sections)}")
# Count paragraphs as rough page estimate
para_count = len(doc.paragraphs)
print(f"Paragraphs: {para_count}")
print(f"Estimated pages: {max(1, para_count // 30)}")
